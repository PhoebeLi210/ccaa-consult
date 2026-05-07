#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智质通·咨询版 - 体系文件生成器

核心功能：
1. 基于企业信息生成ISO体系文件
2. 支持模板+变量替换
3. 支持AI扩写描述性内容
"""

import os
import re
import json
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


@dataclass
class GenerationConfig:
    """生成配置"""
    cover_style: str = "standard"  # simple/standard/business
    has_border: bool = True
    controlled_status: str = "草稿"  # 草稿/受控/作废
    numbering_rule: str = "Q/{company}-{dept}-{year}-{seq}"
    last_audit_date: Optional[str] = None


class DocumentGenerator:
    """体系文件生成器"""

    # ISO9001强制程序文件
    ISO9001_PROCEDURES = [
        "文件控制程序",
        "记录控制程序",
        "内部审核程序",
        "不合格品控制程序",
        "纠正措施程序",
        "预防措施程序",
    ]

    # ISO14001强制程序文件
    ISO14001_PROCEDURES = [
        "环境因素识别与评价控制程序",
        "法律法规及其他要求控制程序",
        "应急准备和响应控制程序",
        "监测和测量控制程序",
        "合规性评价控制程序",
    ]

    # ISO45001强制程序文件
    ISO45001_PROCEDURES = [
        "危险源辨识、风险评价和控制措施确定程序",
        "能力、培训和意识控制程序",
        "沟通和参与控制程序",
        "应急准备和响应控制程序",
        "绩效监测和测量控制程序",
    ]

    def __init__(self, template_dir: str = None, llm_client=None):
        """
        初始化生成器

        Args:
            template_dir: 模板目录
            llm_client: 大模型客户端（用于AI扩写）
        """
        self.template_dir = template_dir or "./templates"
        self.llm_client = llm_client

    def generate_all(
        self,
        company_info: Dict[str, Any],
        standards: List[str],
        config: GenerationConfig = None
    ) -> Dict[str, Any]:
        """
        生成全套体系文件

        Args:
            company_info: 企业信息
            standards: 目标标准列表
            config: 生成配置

        Returns:
            生成的文件列表
        """
        if config is None:
            config = GenerationConfig()

        result = {
            "manuals": [],
            "procedures": [],
            "records": [],
            "instructions": [],
            "forms": [],
        }

        # 1. 生成管理手册
        for standard in standards:
            manual = self._generate_manual(company_info, standard, config)
            result["manuals"].append(manual)

        # 2. 生成程序文件
        for standard in standards:
            procedures = self._generate_procedures(company_info, standard, config)
            result["procedures"].extend(procedures)

        # 3. 生成记录表格
        records = self._generate_records(company_info, standards, config)
        result["records"].extend(records)

        # 4. 生成作业指导书（如果有生产过程）
        if company_info.get("main_processes"):
            instructions = self._generate_instructions(company_info, config)
            result["instructions"].extend(instructions)

        return result

    def _generate_manual(
        self,
        company_info: Dict[str, Any],
        standard: str,
        config: GenerationConfig
    ) -> Dict[str, Any]:
        """生成管理手册"""
        standard_names = {
            "ISO9001": "质量",
            "ISO14001": "环境",
            "ISO45001": "职业健康安全",
        }

        standard_name = standard_names.get(standard, standard)

        # 手册内容模板
        content = f"""
# {company_info.get('company_name', 'XX公司')}{standard_name}管理手册

## 文件信息
- 文件编号：{self._generate_file_number(company_info, config, 'A-001')}
- 版本：A/0
- 生效日期：{datetime.now().strftime('%Y年%m月%d日')}
- 受控状态：{config.controlled_status}

## 目录

### 1. 前言
#### 1.1 公司简介
{self._generate_company_intro(company_info)}

#### 1.2 手册管理
本手册依据{standard}标准编制，描述了公司{standard_name}管理体系的范围、过程和相互作用。

### 2. 范围
#### 2.1 总则
本手册规定了公司{standard_name}管理体系的要求，适用于公司{company_info.get('industry', '相关')}相关活动。

#### 2.2 应用范围
本手册适用于公司{company_info.get('company_name', '')}的{standard_name}管理体系。

### 3. 引用标准
- {standard}标准

### 4. 术语和定义
本手册采用{standard}标准中的术语和定义。

### 5. 组织环境
#### 5.1 理解组织及其环境
{self._generate_context_analysis(company_info, standard)}

#### 5.2 理解相关方的需求和期望
{self._generate_stakeholder_analysis(company_info, standard)}

#### 5.3 确定管理体系范围
本管理体系覆盖公司所有与{standard_name}相关的活动和过程。

#### 5.4 管理体系及其过程
{self._generate_process_description(company_info, standard)}

### 6. 领导作用
#### 6.1 领导作用和承诺
{self._generate_leadership(company_info, standard)}

#### 6.2 方针
{self._generate_policy(company_info, standard)}

#### 6.3 组织的岗位、职责和权限
{self._generate_responsibilities(company_info, standard)}

### 7. 策划
#### 7.1 应对风险和机遇的措施
{self._generate_risk_opportunity(company_info, standard)}

#### 7.2 目标及其实现的策划
{self._generate_objectives(company_info, standard)}

### 8. 支持
#### 8.1 资源
{self._generate_resources(company_info, standard)}

#### 8.2 能力
{self._generate_competence(company_info, standard)}

#### 8.3 意识
{self._generate_awareness(company_info, standard)}

#### 8.4 信息交流
{self._generate_communication(company_info, standard)}

#### 8.5 文件化信息
{self._generate_documentation(company_info, standard)}

### 9. 运行
{self._generate_operation(company_info, standard)}

### 10. 绩效评价
#### 10.1 监视、测量、分析和评价
{self._generate_monitoring(company_info, standard)}

#### 10.2 内部审核
{self._generate_internal_audit(company_info, standard)}

#### 10.3 管理评审
{self._generate_management_review(company_info, standard)}

### 11. 改进
#### 11.1 总则
公司持续改进{standard_name}管理体系的适宜性、充分性和有效性。

#### 11.2 不合格和纠正措施
详见《不合格和纠正措施控制程序》。

#### 11.3 持续改进
公司利用{standard_name}方针、目标、审核结果、数据分析、纠正措施和管理评审，持续改进管理体系的有效性。

## 附录
- 附录A：组织架构图
- 附录B：程序文件清单
- 附录C：记录清单
"""

        return {
            "doc_type": "manual",
            "title": f"{standard_name}管理手册",
            "file_name": f"{standard_name}管理手册.docx",
            "content": content,
            "standard": standard,
        }

    def _generate_procedures(
        self,
        company_info: Dict[str, Any],
        standard: str,
        config: GenerationConfig
    ) -> List[Dict[str, Any]]:
        """生成程序文件"""
        procedures = []

        # 根据标准选择程序文件
        if standard == "ISO9001":
            procedure_list = self.ISO9001_PROCEDURES
        elif standard == "ISO14001":
            procedure_list = self.ISO14001_PROCEDURES
        elif standard == "ISO45001":
            procedure_list = self.ISO45001_PROCEDURES
        else:
            procedure_list = []

        for i, procedure_name in enumerate(procedure_list, 1):
            procedure = self._generate_procedure(company_info, procedure_name, standard, config, i)
            procedures.append(procedure)

        return procedures

    def _generate_procedure(
        self,
        company_info: Dict[str, Any],
        procedure_name: str,
        standard: str,
        config: GenerationConfig,
        seq: int
    ) -> Dict[str, Any]:
        """生成单个程序文件"""
        file_number = self._generate_file_number(company_info, config, f'B-{seq:03d}')

        content = f"""
# {procedure_name}

## 文件信息
- 文件编号：{file_number}
- 版本：A/0
- 生效日期：{datetime.now().strftime('%Y年%m月%d日')}
- 受控状态：{config.controlled_status}

## 1. 目的
规定{procedure_name.replace('控制程序', '').replace('程序', '')}的要求，确保体系有效运行。

## 2. 范围
本程序适用于公司所有相关活动和过程。

## 3. 职责
### 3.1 总经理
- 批准本程序

### 3.2 管理者代表
- 组织实施本程序

### 3.3 各部门
- 按本程序要求执行相关工作

## 4. 工作程序
### 4.1 流程图
（待补充）

### 4.2 详细说明
（待补充具体操作步骤）

## 5. 相关文件
- 管理手册

## 6. 相关记录
- 相关记录表格
"""

        return {
            "doc_type": "procedure",
            "title": procedure_name,
            "file_name": f"{procedure_name}.docx",
            "content": content,
            "standard": standard,
        }

    def _generate_records(
        self,
        company_info: Dict[str, Any],
        standards: List[str],
        config: GenerationConfig
    ) -> List[Dict[str, Any]]:
        """生成记录表格"""
        records = []

        # 通用记录
        common_records = [
            "文件发放回收记录",
            "记录控制一览表",
            "培训记录表",
            "内部审核计划",
            "内部审核检查表",
            "内部审核报告",
            "管理评审计划",
            "管理评审报告",
            "纠正措施记录",
            "预防措施记录",
        ]

        for record_name in common_records:
            records.append({
                "doc_type": "record",
                "title": record_name,
                "file_name": f"{record_name}.docx",
                "content": self._generate_record_template(record_name, company_info),
            })

        # 环境相关记录
        if "ISO14001" in standards:
            env_records = [
                "环境因素识别与评价表",
                "重要环境因素清单",
                "法律法规清单",
                "环境运行检查记录",
                "应急演练记录",
            ]
            for record_name in env_records:
                records.append({
                    "doc_type": "record",
                    "title": record_name,
                    "file_name": f"{record_name}.docx",
                    "content": self._generate_record_template(record_name, company_info),
                })

        # 职业健康安全相关记录
        if "ISO45001" in standards:
            ohs_records = [
                "危险源辨识与评价表",
                "不可接受风险清单",
                "安全检查记录",
                "劳保用品发放记录",
                "职业健康体检记录",
            ]
            for record_name in ohs_records:
                records.append({
                    "doc_type": "record",
                    "title": record_name,
                    "file_name": f"{record_name}.docx",
                    "content": self._generate_record_template(record_name, company_info),
                })

        return records

    def _generate_instructions(
        self,
        company_info: Dict[str, Any],
        config: GenerationConfig
    ) -> List[Dict[str, Any]]:
        """生成作业指导书"""
        instructions = []

        for process in company_info.get("main_processes", []):
            instruction = {
                "doc_type": "instruction",
                "title": f"{process}作业指导书",
                "file_name": f"{process}作业指导书.docx",
                "content": self._generate_instruction_template(process, company_info),
            }
            instructions.append(instruction)

        return instructions

    def _generate_instruction_template(
        self,
        process: str,
        company_info: Dict[str, Any]
    ) -> str:
        """生成作业指导书模板"""
        return f"""
# {process}作业指导书

## 文件信息
- 文件编号：WI-XXX-XXX
- 版本：A/0
- 生效日期：{datetime.now().strftime('%Y年%m月%d日')}

## 1. 目的
规范{process}作业过程，确保产品质量。

## 2. 适用范围
适用于{process}工序的操作。

## 3. 职责
- 操作工：按本指导书进行操作
- 班组长：监督执行情况
- 质检员：检验产品质量

## 4. 操作步骤
### 4.1 准备工作
1. 检查设备状态
2. 准备所需工具和材料
3. 穿戴好劳保用品

### 4.2 操作步骤
（待补充具体操作步骤）

### 4.3 注意事项
（待补充安全注意事项）

## 5. 质量要求
（待补充产品质量要求）

## 6. 相关记录
- 生产记录表
- 设备点检表
"""

    def _generate_record_template(self, record_name: str, company_info: Dict[str, Any]) -> str:
        """生成记录表格模板"""
        return f"""
# {record_name}

## 基本信息
- 编号：XXX-XXX
- 日期：____年____月____日
- 部门：____________

## 记录内容
| 序号 | 内容 | 备注 |
|------|------|------|
| 1 | | |
| 2 | | |
| 3 | | |

## 签字确认
- 编制：____________ 日期：____________
- 审核：____________ 日期：____________
- 批准：____________ 日期：____________
"""

    def _generate_file_number(
        self,
        company_info: Dict[str, Any],
        config: GenerationConfig,
        file_code: str
    ) -> str:
        """生成文件编号"""
        company_abbr = company_info.get("company_name", "XX")[:2].upper() if company_info.get("company_name") else "XX"
        year = datetime.now().year
        return f"Q/{company_abbr}-{file_code}-{year}"

    def _generate_company_intro(self, company_info: Dict[str, Any]) -> str:
        """生成公司简介"""
        name = company_info.get("company_name", "本公司")
        industry = company_info.get("industry", "")
        employees = company_info.get("employee_count", "")
        area = company_info.get("office_area_sqm", "")

        intro = f"{name}是一家{industry}企业"
        if employees:
            intro += f"，现有员工{employees}人"
        if area:
            intro += f"，办公面积{area}平方米"
        intro += "。"

        return intro

    def _generate_context_analysis(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成组织环境分析"""
        return f"""
公司对内外部环境进行了分析，识别了以下关键因素：

**内部环境：**
- 组织结构：{', '.join(company_info.get('departments', ['各部门']))}
- 资源状况：员工{company_info.get('employee_count', '若干')}人，主要设备{len(company_info.get('main_equipment', []))}类
- 企业文化：追求卓越，持续改进

**外部环境：**
- 市场竞争：行业竞争激烈，需不断提升服务质量
- 法律法规：需遵守相关法律法规要求
- 技术发展：关注行业技术发展趋势
"""

    def _generate_stakeholder_analysis(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成相关方分析"""
        return """
公司识别了以下相关方及其需求：

| 相关方 | 需求和期望 |
|--------|-----------|
| 客户 | 满意的产品和服务 |
| 员工 | 良好的工作环境和发展机会 |
| 供应商 | 公平合作、及时付款 |
| 政府部门 | 合规经营 |
| 社区 | 环境保护、社会责任 |
"""

    def _generate_process_description(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成过程描述"""
        processes = company_info.get("main_processes", [])
        if processes:
            return f"公司主要过程包括：{'、'.join(processes)}。各过程按照标准要求进行控制。"
        return "公司按照标准要求建立了管理体系过程。"

    def _generate_leadership(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成领导作用描述"""
        return """
公司最高管理者通过以下方式证实其领导作用和承诺：
1. 确保管理体系方针和目标得到建立，并与组织的战略方向相适应
2. 确保管理体系要求融入组织的业务过程
3. 促进使用过程方法和基于风险的思维
4. 确保管理体系所需资源的获得
5. 沟通有效的管理体系和符合管理体系要求的重要性
6. 确保管理体系实现其预期的结果
7. 指导和支持员工为管理体系的有效性做出贡献
8. 促进持续改进
9. 支持其他相关管理人员在其职责范围内证实其领导作用
"""

    def _generate_policy(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成方针"""
        standard_names = {
            "ISO9001": "质量",
            "ISO14001": "环境",
            "ISO45001": "职业健康安全",
        }
        name = standard_names.get(standard, "")

        return f"""
## {name}方针

**{company_info.get('company_name', '本公司')}{name}方针：**

遵纪守法  顾客至上
持续改进  追求卓越

**方针释义：**
- 遵纪守法：严格遵守国家法律法规和行业标准要求
- 顾客至上：以顾客为关注焦点，满足顾客需求和期望
- 持续改进：不断完善管理体系，提高管理水平
- 追求卓越：追求卓越绩效，实现可持续发展
"""

    def _generate_responsibilities(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成职责权限"""
        departments = company_info.get("departments", [])
        dept_text = "、".join(departments) if departments else "各部门"

        return f"""
公司建立了清晰的组织架构，明确了各部门的职责和权限：

**总经理：**
- 全面负责公司经营管理
- 批准管理手册和程序文件
- 任命管理者代表

**管理者代表：**
- 建立和保持管理体系
- 向最高管理者报告管理体系绩效
- 提高员工满足顾客和法规要求的意识

**{dept_text}：**
- 按职责分工执行管理体系要求
- 完成部门目标和指标
- 参与持续改进活动
"""

    def _generate_risk_opportunity(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成风险和机遇"""
        return """
公司建立了风险和机遇管理机制：

1. 风险识别：识别可能影响管理体系预期结果的风险
2. 风险评价：评估风险的严重程度和发生可能性
3. 风险应对：制定并实施风险应对措施
4. 机遇把握：识别并利用改进的机会

详见《风险和机遇应对措施控制程序》。
"""

    def _generate_objectives(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成目标"""
        quality_goals = company_info.get("quality_goals", "产品合格率≥95%")

        return f"""
公司制定了以下管理目标：

**质量目标：**
- {quality_goals}
- 客户满意度≥90%

**环境目标：**
- 固体废弃物合规处置率100%
- 节能降耗达标

**职业健康安全目标：**
- 火灾事故为0
- 人身伤害事故为0

目标按年度分解到各部门，定期考核。
"""

    def _generate_resources(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成资源描述"""
        equipment = company_info.get("main_equipment", [])
        equipment_text = "、".join(equipment) if equipment else "各类设备"

        return f"""
公司确定并提供了管理体系所需的资源：

**人力资源：**
- 员工{company_info.get('employee_count', '若干')}人
- 各岗位人员具备相应能力

**基础设施：**
- 办公面积{company_info.get('office_area_sqm', '若干')}平方米
- 主要设备：{equipment_text}

**过程运行环境：**
- 提供适宜的工作环境
- 确保安全生产条件

**监视和测量资源：**
- 配备必要的检测设备
- 定期校准和维护
"""

    def _generate_competence(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成能力描述"""
        return """
公司建立了人员能力管理机制：

1. 确定各岗位所需的能力要求
2. 评估现有人员的能力水平
3. 制定培训计划，弥补能力差距
4. 评价培训效果，确保能力提升

详见《人力资源控制程序》。
"""

    def _generate_awareness(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成意识描述"""
        return """
公司通过培训、会议、文件传达等方式，确保员工了解：

1. 管理方针
2. 管理目标
3. 自身对管理体系有效性的贡献
4. 不符合管理体系要求的后果
"""

    def _generate_communication(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成信息交流描述"""
        return """
公司建立了内外部信息交流机制：

**内部交流：**
- 部门例会
- 管理评审会议
- 内部审核
- 文件传达

**外部交流：**
- 与客户的沟通
- 与供应商的沟通
- 与政府部门的沟通
- 与其他相关方的沟通

详见《信息交流和沟通控制程序》。
"""

    def _generate_documentation(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成文件化信息描述"""
        return """
公司管理体系的文件化信息包括：

**一级文件：**
- 管理手册

**二级文件：**
- 程序文件

**三级文件：**
- 作业指导书
- 管理制度

**四级文件：**
- 记录表格

详见《文件控制程序》和《记录控制程序》。
"""

    def _generate_operation(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成运行描述"""
        processes = company_info.get("main_processes", [])
        processes_text = "、".join(processes) if processes else "各生产过程"

        return f"""
公司对运行过程进行策划和控制：

**运行策划：**
- 确定产品和服务的要求
- 建立过程控制准则
- 配备所需资源

**运行控制：**
- 主要过程：{processes_text}
- 按作业指导书进行操作
- 做好过程记录

**产品和服务放行：**
- 按检验规范进行检验
- 合格后方可放行

**不合格输出控制：**
- 识别和控制不合格品
- 采取纠正措施
"""

    def _generate_monitoring(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成监视测量描述"""
        return """
公司对管理体系进行监视和测量：

**顾客满意：**
- 定期进行顾客满意度调查
- 分析顾客反馈信息

**内部审核：**
- 按计划进行内部审核
- 验证管理体系符合性和有效性

**过程监视：**
- 监视关键过程参数
- 分析过程绩效

**产品监视：**
- 检验产品质量
- 控制不合格品
"""

    def _generate_internal_audit(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成内部审核描述"""
        return """
公司按计划进行内部审核：

1. 编制年度内审计划
2. 组建审核组
3. 实施现场审核
4. 编制审核报告
5. 跟踪验证纠正措施

详见《内部审核控制程序》。
"""

    def _generate_management_review(self, company_info: Dict[str, Any], standard: str) -> str:
        """生成管理评审描述"""
        return """
公司定期进行管理评审：

**评审输入：**
- 审核结果
- 顾客反馈
- 过程绩效
- 不合格情况
- 纠正措施状态
- 以往评审的跟踪措施
- 变更信息
- 改进建议

**评审输出：**
- 改进决定
- 资源需求
- 管理体系变更

详见《管理评审控制程序》。
"""

    def create_docx(self, content: str, output_path: str, title: str = None) -> str:
        """
        创建Word文档

        Args:
            content: 文档内容
            output_path: 输出路径
            title: 文档标题

        Returns:
            生成的文件路径
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)

        # 添加标题
        if title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.size = Pt(18)
            run.bold = True

        # 添加内容
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                # 一级标题
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # 二级标题
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # 三级标题
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                # 四级标题
                p = doc.add_heading(line[5:], level=4)
            elif line.startswith('**') and line.endswith('**'):
                # 加粗段落
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('| '):
                # 表格行（简化处理）
                p = doc.add_paragraph(line)
            elif line.startswith('- '):
                # 列表项
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                # 普通段落
                p = doc.add_paragraph(line)

        # 保存文档
        doc.save(output_path)
        return output_path


# 便捷函数
def generate_documents(
    company_info: Dict[str, Any],
    standards: List[str],
    output_dir: str = None,
    config: GenerationConfig = None
) -> Dict[str, Any]:
    """
    生成体系文件的便捷函数

    Args:
        company_info: 企业信息
        standards: 目标标准列表
        output_dir: 输出目录
        config: 生成配置

    Returns:
        生成的文件列表
    """
    generator = DocumentGenerator()
    return generator.generate_all(company_info, standards, config)
